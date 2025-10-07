# app.py
import streamlit as st
import altair as alt
import plotly.express as px
import pandas as pd
import numpy as np
from datetime import datetime
import joblib
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from PyPDF2 import PdfReader
import requests
import base64
import json
import urllib.parse

# Tracking utils
from track_utils import (
    create_page_visited_table,
    add_page_visited_details,
    view_all_page_visited_details,
    add_prediction_details,
    view_all_prediction_details,
    create_emotionclf_table,
    IST
)

# ---------- Load model ----------
try:
    pipe_lr = joblib.load(open("./models/emotion_classifier_pipe_lr.pkl", "rb"))
except Exception as e:
    st = getattr(__import__("streamlit"), "st")
    st.error(f"Failed to load model: {e}")
    pipe_lr = None

# ---------- Prediction helpers ----------
def predict_emotions(docx):
    if pipe_lr is None:
        return "error"
    results = pipe_lr.predict([docx])
    return results[0]

def get_prediction_proba(docx):
    if pipe_lr is None:
        return np.array([[1.0]])
    results = pipe_lr.predict_proba([docx])
    return results

# ---------- Emoji dictionary ----------
emotions_emoji_dict = {
    "anger": "😠", "disgust": "🪮", "fear": "😨😱", "happy": "🫷",
    "joy": "😂", "neutral": "😐", "sad": "😔", "sadness": "😔",
    "shame": "😳", "surprise": "😮"
}

# ---------- File parsing helpers ----------
def extract_text_from_pdf(file_bytes):
    text_lines = []
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            page_lines = [ln.strip() for ln in page_text.splitlines() if ln.strip()]
            text_lines.extend(page_lines)
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
    return text_lines

def extract_text_from_docx(file_bytes):
    text_lines = []
    try:
        temp = io.BytesIO(file_bytes)
        doc = Document(temp)
        for para in doc.paragraphs:
            line = para.text.strip()
            if line:
                text_lines.append(line)
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
    return text_lines

def extract_text_from_txt(file_bytes):
    text = file_bytes.decode('utf-8', errors='replace')
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    return lines

def extract_lines_from_csv_or_excel(file_bytes, file_type, chosen_column=None):
    try:
        if file_type == 'csv':
            df = pd.read_csv(io.BytesIO(file_bytes))
        else:
            df = pd.read_excel(io.BytesIO(file_bytes))
    except Exception as e:
        st.error(f"Error reading {file_type.upper()}: {e}")
        return []

    if df.empty:
        return []

    if chosen_column and chosen_column in df.columns:
        lines = df[chosen_column].astype(str).fillna("").tolist()
    else:
        lines = df.astype(str).fillna("").apply(lambda r: " | ".join([c for c in r.values if str(c).strip()]), axis=1).tolist()

    lines = [ln for ln in lines if ln and str(ln).strip()]
    return lines

# ---------- Export helpers ----------
def dataframe_to_csv_bytes(df):
    return df.to_csv(index=False).encode('utf-8')

def dataframe_to_excel_bytes(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='EmotionResults')
    output.seek(0)
    return output.getvalue()

def dataframe_to_txt_bytes(df):
    buff = io.StringIO()
    for idx, row in df.iterrows():
        line_num = int(row['Line']) if 'Line' in row and pd.notna(row['Line']) else idx + 1
        text = row.get('Text', '')
        emotion = row.get('Predicted Emotion', '')
        emoji = row.get('Emoji', '')
        confidence = row.get('Confidence', 0)
        buff.write(f"Line {line_num}: {text}\n")
        buff.write(f"Prediction: {emotion} {emoji}\n")
        buff.write(f"Confidence: {confidence}%\n")
        buff.write("-" * 60 + "\n")
    return buff.getvalue().encode('utf-8')

def dataframe_to_pdf_bytes(df, title="Emotion Detection Report"):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, height - 72, title)
    c.setFont("Helvetica", 10)
    y = height - 100
    line_height = 12
    for idx, row in df.iterrows():
        line_num = int(row['Line']) if 'Line' in row and pd.notna(row['Line']) else idx + 1
        text_line = f"Line {line_num}: {row.get('Text','')}"
        emotion_line = f"Prediction: {row.get('Predicted Emotion','')} {row.get('Emoji','')} - Confidence: {row.get('Confidence',0)}%"
        for text in [text_line, emotion_line]:
            while text:
                max_chars = 90
                part = text[:max_chars]
                c.drawString(72, y, part)
                text = text[max_chars:]
                y -= line_height
                if y < 72:
                    c.showPage()
                    c.setFont("Helvetica", 10)
                    y = height - 72
        y -= 6
        if y < 72:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - 72
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def dataframe_to_docx_bytes(df):
    doc = Document()
    doc.add_heading('Emotion Detection Report', level=1)
    for idx, row in df.iterrows():
        line_num = int(row['Line']) if 'Line' in row and pd.notna(row['Line']) else idx + 1
        text = row.get('Text', '')
        emotion = row.get('Predicted Emotion', '')
        emoji = row.get('Emoji', '')
        confidence = row.get('Confidence', 0)
        p = doc.add_paragraph()
        p.add_run(f"Line {line_num}: ").bold = True
        p.add_run(str(text))
        doc.add_paragraph(f"Prediction: {emotion} {emoji}")
        doc.add_paragraph(f"Confidence: {confidence}%")
        doc.add_paragraph("-" * 40)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

# ---------- Extension / Query Param helpers ----------
def decode_b64_json_param(b64_str):
    try:
        b64_decoded = urllib.parse.unquote(b64_str)
        json_bytes = base64.b64decode(b64_decoded)
        obj = json.loads(json_bytes.decode('utf-8'))
        return obj
    except Exception:
        try:
            return json.loads(urllib.parse.unquote(b64_str))
        except Exception:
            return None

def fetch_payload_from_receiver(token, receiver_base):
    try:
        url = f"{receiver_base.rstrip('/')}/payload/{token}"
        resp = requests.get(url, timeout=10)
        resp.raise_for_status()
        return resp.json()
    except Exception as e:
        st.error(f"Failed to fetch payload from receiver: {e}")
        return None

# ---------- Payload processing ----------
def process_payload_and_show(payload, source_name="Extension Payload"):
    lines = []
    if isinstance(payload.get('reviews'), (list, tuple)) and payload.get('reviews'):
        lines = [str(r).strip() for r in payload.get('reviews') if str(r).strip()]
    else:
        for k in ('title', 'description', 'text', 'product_description'):
            v = payload.get(k)
            if v:
                if isinstance(v, (list, tuple)):
                    lines.extend([str(x).strip() for x in v if str(x).strip()])
                else:
                    lines.append(str(v).strip())
        if isinstance(payload, (list, tuple)):
            lines.extend([str(x).strip() for x in payload if str(x).strip()])

    lines = [ln for i, ln in enumerate(lines) if ln and ln not in lines[:i]]

    if not lines:
        st.info("No textual entries found in the incoming payload.")
        return

    st.markdown(f"## Results from: {source_name}")
    if payload.get('title'):
        st.markdown(f"**Title:** {payload.get('title')}")
    if payload.get('url'):
        st.markdown(f"[Open product URL]({payload.get('url')})")
    if payload.get('images'):
        st.markdown("#### Product images")
        imgs = payload.get('images')[:8]
        cols = st.columns(min(4, len(imgs)))
        for i, img_url in enumerate(imgs):
            try:
                with cols[i % len(cols)]:
                    st.image(img_url, width=150)
            except Exception:
                pass

    st.success(f"Processing {len(lines)} text items from {source_name}")
    results = []
    progress_bar = st.progress(0)
    for i, line in enumerate(lines, start=1):
        text_for_model = str(line)[:4000]
        pred = predict_emotions(text_for_model)
        proba = get_prediction_proba(text_for_model)
        conf = round(np.max(proba) * 100, 2) if proba is not None else 0.0
        emoji = emotions_emoji_dict.get(pred, "")
        results.append({
            "Line": i,
            "Text": line,
            "Predicted Emotion": pred,
            "Confidence": conf,
            "Emoji": emoji
        })
        try:
            add_prediction_details(line, pred, np.max(proba), datetime.now(IST))
        except Exception:
            pass
        progress_bar.progress(int(i / len(lines) * 100))

    df_results = pd.DataFrame(results)
    st.markdown("### Line-by-line Emotion Results (From Extension)")
    st.dataframe(df_results)

    st.markdown("### Summary")
    pred_count = df_results['Predicted Emotion'].value_counts().rename_axis('Prediction').reset_index(name='Counts')
    if not pred_count.empty:
        bar_chart = alt.Chart(pred_count).mark_bar().encode(
            x=alt.X('Prediction', sort='-y'),
            y='Counts',
            color='Prediction'
        )
        st.altair_chart(bar_chart, use_container_width=True)
        pie_chart = px.pie(pred_count, values='Counts', names='Prediction', title='Emotion Distribution (From Extension)')
        st.plotly_chart(pie_chart, use_container_width=True)

    st.markdown("### Sample lines with emojis")
    sample = df_results.head(10)
    for _, r in sample.iterrows():
        st.write(f"**Line {int(r['Line'])}** — {r['Emoji']}  {r['Text']}")
        st.caption(f"{r['Predicted Emotion'].capitalize()} — Confidence: {r['Confidence']}%")

    st.markdown("---")
    st.subheader("📤 Download Results")
    st.download_button("Download CSV", data=dataframe_to_csv_bytes(df_results),
                       file_name="extension_emotion_results.csv", mime='text/csv')
    st.download_button("Download XLSX", data=dataframe_to_excel_bytes(df_results),
                       file_name="extension_emotion_results.xlsx",
                       mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    st.download_button("Download TXT", data=dataframe_to_txt_bytes(df_results),
                       file_name="extension_emotion_results.txt", mime='text/plain')
    st.download_button("Download PDF", data=dataframe_to_pdf_bytes(df_results, title=f"Emotion Analysis - {source_name}"),
                       file_name="extension_emotion_results.pdf", mime='application/pdf')
    st.download_button("Download DOCX", data=dataframe_to_docx_bytes(df_results),
                       file_name="extension_emotion_results.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ---------- Main App ----------
def main():
    st.set_page_config(page_title="Emotion Detector", page_icon="💬", layout="centered")
    st.markdown("<h1 style='text-align: center; color: #6c63ff;'>💬 Emotion Detection in Text & Files</h1>", unsafe_allow_html=True)
    st.markdown("<h4 style='text-align: center; color: grey;'>Analyze emotions line-by-line from many file types.</h4>", unsafe_allow_html=True)
    st.write("")

    create_page_visited_table()
    create_emotionclf_table()

    params = st.experimental_get_query_params()
    if 'data' in params and params.get('data'):
        st.info("Detected incoming data parameter. Decoding...")
        payload = decode_b64_json_param(params['data'][0])
        if payload is None:
            st.error("Could not decode `data` query parameter.")
        else:
            process_payload_and_show(payload, source_name="Direct (data param)")
            return

    if 'token' in params and 'receiver' in params and params.get('token') and params.get('receiver'):
        st.info("Detected token + receiver query parameters. Fetching payload...")
        token = params['token'][0]
        receiver_base = params['receiver'][0]
        payload = fetch_payload_from_receiver(token, receiver_base)
        if payload:
            process_payload_and_show(payload, source_name=f"Receiver ({receiver_base})")
            return
        else:
            st.error("Failed to fetch payload for token from receiver.")

    menu = ["Home", "File Analysis", "Monitor", "About"]
    choice = st.sidebar.selectbox("Menu", menu)

    if choice == "Home":
        add_page_visited_details("Home", datetime.now(IST))
        st.subheader("Try typing some text below:")
        with st.form(key='emotion_clf_form'):
            raw_text = st.text_area("Your Text", height=150, placeholder="Enter your message here...")
            submit_text = st.form_submit_button(label='Analyze Emotion')

        if submit_text and raw_text.strip():
            col1, col2 = st.columns(2)
            prediction = predict_emotions(raw_text)
            probability = get_prediction_proba(raw_text)
            confidence = round(np.max(probability)*100, 2) if probability is not None else 0.0
            add_prediction_details(raw_text, prediction, np.max(probability), datetime.now(IST))

            with col1:
                st.success("Original Text")
                st.write(raw_text)
                st.success("Prediction")
                emoji_icon = emotions_emoji_dict.get(prediction, "")
                st.write(f"{prediction.capitalize()} {emoji_icon}")
                st.markdown(f"**Confidence:** {confidence}%")

            with col2:
                st.success("Prediction Probability")
                proba_df = pd.DataFrame(probability, columns=pipe_lr.classes_) if pipe_lr is not None else pd.DataFrame()
                if not proba_df.empty:
                    proba_df_clean = proba_df.T.reset_index()
                    proba_df_clean.columns = ["Emotions", "Probability"]
                    fig = alt.Chart(proba_df_clean).mark_bar().encode(
                        x=alt.X('Emotions', sort='-y'),
                        y='Probability',
                        color='Emotions'
                    )
                    st.altair_chart(fig, use_container_width=True)
                else:
                    st.write("Probability data not available.")

            st.markdown("---")
            st.subheader("📤 Export Results")
            csv_data = pd.DataFrame({
                "Text": [raw_text],
                "Predicted Emotion": [prediction],
                "Confidence": [confidence]
            })
            csv = dataframe_to_csv_bytes(csv_data)
            st.download_button("Download as CSV", data=csv, file_name="emotion_result.csv", mime='text/csv')
            pdf_buffer = dataframe_to_pdf_bytes(csv_data)
            st.download_button("Download as PDF", data=pdf_buffer, file_name="emotion_report.pdf", mime='application/pdf')

    # ---------- File Analysis ----------
    elif choice == "File Analysis":
        add_page_visited_details("File Analysis", datetime.now(IST))
        st.subheader("Upload a file to analyze (txt, csv, xls/xlsx, pdf, docx)")
        uploaded_file = st.file_uploader("Choose a file", type=['txt','csv','xls','xlsx','pdf','docx'])

        if uploaded_file:
            file_bytes = uploaded_file.read()
            file_name = uploaded_file.name.lower()
            chosen_column = None
            lines = []

            if file_name.endswith('.pdf'):
                with st.spinner("Extracting text from PDF..."):
                    lines = extract_text_from_pdf(file_bytes)
            elif file_name.endswith('.docx'):
                with st.spinner("Extracting text from DOCX..."):
                    lines = extract_text_from_docx(file_bytes)
            elif file_name.endswith('.txt'):
                with st.spinner("Reading TXT..."):
                    lines = extract_text_from_txt(file_bytes)
            elif file_name.endswith('.csv') or file_name.endswith('.xls') or file_name.endswith('.xlsx'):
                try:
                    if file_name.endswith('.csv'):
                        preview_df = pd.read_csv(io.BytesIO(file_bytes), nrows=200)
                        file_type = 'csv'
                    else:
                        preview_df = pd.read_excel(io.BytesIO(file_bytes), nrows=200)
                        file_type = 'excel'
                except Exception as e:
                    st.error(f"Could not read file for preview: {e}")
                    preview_df = pd.DataFrame()
                    file_type = 'csv' if file_name.endswith('.csv') else 'excel'

                if not preview_df.empty:
                    st.markdown("### File Preview (first rows)")
                    st.dataframe(preview_df.head(10))
                    col_options = ["<Combine all columns>"] + list(preview_df.columns.astype(str))
                    chosen_column = st.selectbox("Select a specific column to analyze (or choose combine)", col_options)
                    if chosen_column == "<Combine all columns>":
                        chosen_column = None

                with st.spinner("Extracting lines from table..."):
                    lines = extract_lines_from_csv_or_excel(file_bytes, 'csv' if file_name.endswith('.csv') else 'excel', chosen_column)

            if not lines:
                st.warning("No text found in the uploaded file.")
                return

            st.info(f"Found {len(lines)} lines. Predicting emotions...")
            results = []
            progress_bar = st.progress(0)
            for i, line in enumerate(lines, start=1):
                text_for_model = str(line)[:4000]
                pred = predict_emotions(text_for_model)
                proba = get_prediction_proba(text_for_model)
                conf = round(np.max(proba)*100, 2) if proba is not None else 0.0
                emoji = emotions_emoji_dict.get(pred, "")
                results.append({
                    "Line": i,
                    "Text": line,
                    "Predicted Emotion": pred,
                    "Confidence": conf,
                    "Emoji": emoji
                })
                progress_bar.progress(int(i / len(lines) * 100))
            df_results = pd.DataFrame(results)

            st.markdown("### File Analysis Results")
            st.dataframe(df_results)

            st.subheader("📤 Download Results")
            st.download_button("Download CSV", data=dataframe_to_csv_bytes(df_results),
                               file_name="file_emotion_results.csv", mime='text/csv')
            st.download_button("Download XLSX", data=dataframe_to_excel_bytes(df_results),
                               file_name="file_emotion_results.xlsx",
                               mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            st.download_button("Download TXT", data=dataframe_to_txt_bytes(df_results),
                               file_name="file_emotion_results.txt", mime='text/plain')
            st.download_button("Download PDF", data=dataframe_to_pdf_bytes(df_results, title=f"Emotion Analysis - {uploaded_file.name}"),
                               file_name="file_emotion_results.pdf", mime='application/pdf')
            st.download_button("Download DOCX", data=dataframe_to_docx_bytes(df_results),
                               file_name="file_emotion_results.docx",
                               mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # ---------- Monitor ----------
    elif choice == "Monitor":
        add_page_visited_details("Monitor", datetime.now(IST))
        st.subheader("Past Predictions Log")
        try:
            logs = view_all_prediction_details()
            if logs:
                st.dataframe(pd.DataFrame(logs))
            else:
                st.info("No prediction logs available.")
        except Exception as e:
            st.error(f"Error fetching logs: {e}")

    # ---------- About ----------
    elif choice == "About":
        add_page_visited_details("About", datetime.now(IST))
        st.subheader("About This App")
        st.markdown("""
        **Emotion Detection App** analyzes emotions line-by-line from raw text or uploaded files.
        Uses a trained machine learning pipeline and provides visualizations and downloads in multiple formats.
        """)

if __name__ == "__main__":
    main()
